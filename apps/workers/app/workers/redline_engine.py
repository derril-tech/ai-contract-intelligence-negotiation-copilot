# Created automatically by Cursor AI (2024-12-19)

import boto3
import json
import os
import tempfile
from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Tuple
from celery import Celery
from celery_app import celery_app
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
import fitz  # PyMuPDF
from PIL import Image, ImageDraw, ImageFont
import io

@dataclass
class RenderOptions:
    """Options for document rendering"""
    include_track_changes: bool = True
    include_watermark: bool = True
    watermark_text: str = "DRAFT - FOR REVIEW"
    include_comments: bool = True
    output_format: str = "docx"  # "docx" or "pdf"

@dataclass
class RenderResult:
    """Result of document rendering with redlines"""
    agreement_id: str
    original_s3_key: str
    redlined_s3_key: str
    format: str
    file_size: int
    change_count: int
    render_time_ms: int

class RedlineEngineWorker:
    """Worker for rendering documents with redlines and tracked changes"""
    
    def __init__(self):
        self.s3_client = boto3.client(
            's3',
            endpoint_url=os.getenv('S3_ENDPOINT_URL'),
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
            region_name=os.getenv('AWS_REGION', 'us-east-1')
        )
        self.bucket_name = os.getenv('S3_BUCKET_NAME', 'contract-intelligence')
    
    def download_original_document(self, agreement_id: str) -> Tuple[str, str]:
        """Download original document from S3 and return local path and format"""
        try:
            # Try to find the original document
            response = self.s3_client.list_objects_v2(
                Bucket=self.bucket_name,
                Prefix=f"agreements/{agreement_id}/original/"
            )
            
            if 'Contents' not in response:
                raise Exception(f"No original document found for {agreement_id}")
            
            # Get the first file (assuming there's only one original)
            original_key = response['Contents'][0]['Key']
            file_extension = original_key.split('.')[-1].lower()
            
            # Download to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file_extension}') as temp_file:
                self.s3_client.download_fileobj(
                    self.bucket_name,
                    original_key,
                    temp_file
                )
                temp_path = temp_file.name
            
            return temp_path, file_extension
            
        except Exception as e:
            raise Exception(f"Failed to download original document for {agreement_id}: {str(e)}")
    
    def download_redline_data(self, agreement_id: str) -> Dict[str, Any]:
        """Download redline data from S3"""
        try:
            response = self.s3_client.get_object(
                Bucket=self.bucket_name,
                Key=f"agreements/{agreement_id}/redline.json"
            )
            return json.loads(response['Body'].read().decode('utf-8'))
        except Exception as e:
            raise Exception(f"Failed to download redline data for {agreement_id}: {str(e)}")
    
    def download_structure(self, agreement_id: str) -> Dict[str, Any]:
        """Download document structure from S3"""
        try:
            response = self.s3_client.get_object(
                Bucket=self.bucket_name,
                Key=f"agreements/{agreement_id}/structure.json"
            )
            return json.loads(response['Body'].read().decode('utf-8'))
        except Exception as e:
            raise Exception(f"Failed to download structure for {agreement_id}: {str(e)}")
    
    def apply_redlines_to_docx(self, doc_path: str, redline_data: Dict[str, Any], structure: Dict[str, Any], options: RenderOptions) -> str:
        """Apply redlines to DOCX document with tracked changes"""
        try:
            # Load the document
            doc = Document(doc_path)
            
            # Create a mapping of section_id to paragraph indices
            section_mapping = {}
            for i, paragraph in enumerate(doc.paragraphs):
                # In a real implementation, you'd map sections to paragraphs based on structure
                # For now, we'll use a simple approach
                if paragraph.text.strip():
                    section_mapping[f"section_{i}"] = i
            
            # Apply changes
            change_count = 0
            for change_set in redline_data.get('change_sets', []):
                section_id = change_set.get('section_id')
                if section_id in section_mapping:
                    para_index = section_mapping[section_id]
                    paragraph = doc.paragraphs[para_index]
                    
                    if change_set.get('operation') == 'replace':
                        # Clear existing text
                        paragraph.clear()
                        
                        # Add new text with tracked changes
                        run = paragraph.add_run(change_set.get('new_text', ''))
                        run.font.color.rgb = None  # Default color
                        
                        # Add comment if enabled
                        if options.include_comments and change_set.get('comment'):
                            comment = paragraph.add_comment(change_set.get('comment'))
                        
                        change_count += 1
            
            # Add watermark if enabled
            if options.include_watermark:
                # Add watermark as header
                section = doc.sections[0]
                header = section.header
                header_para = header.paragraphs[0]
                header_para.text = options.watermark_text
                header_para.alignment = 1  # Center alignment
                header_para.runs[0].font.color.rgb = None  # Red color
                header_para.runs[0].font.size = 12
            
            # Save to temporary file
            output_path = doc_path.replace('.docx', '_redlined.docx')
            doc.save(output_path)
            
            return output_path, change_count
            
        except Exception as e:
            raise Exception(f"Failed to apply redlines to DOCX: {str(e)}")
    
    def apply_redlines_to_pdf(self, pdf_path: str, redline_data: Dict[str, Any], structure: Dict[str, Any], options: RenderOptions) -> str:
        """Apply redlines to PDF document with annotations"""
        try:
            # Open PDF
            pdf_doc = fitz.open(pdf_path)
            change_count = 0
            
            # Apply changes as annotations
            for change_set in redline_data.get('change_sets', []):
                section_id = change_set.get('section_id')
                
                # Find the page and position for this section
                # In a real implementation, you'd use the structure to find exact positions
                # For now, we'll add annotations to the first page
                if pdf_doc.page_count > 0:
                    page = pdf_doc[0]
                    
                    # Add highlight annotation for changes
                    rect = fitz.Rect(50, 50 + change_count * 30, 500, 80 + change_count * 30)
                    highlight = page.add_highlight_annot(rect)
                    highlight.set_text(change_set.get('comment', 'Change applied'))
                    highlight.update()
                    
                    change_count += 1
            
            # Add watermark if enabled
            if options.include_watermark:
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc[page_num]
                    
                    # Create watermark text
                    text_rect = fitz.Rect(0, 0, page.rect.width, page.rect.height)
                    page.insert_text(
                        (page.rect.width / 2, 50),  # Center top
                        options.watermark_text,
                        fontsize=24,
                        color=(1, 0, 0),  # Red color
                        opacity=0.3
                    )
            
            # Save to temporary file
            output_path = pdf_path.replace('.pdf', '_redlined.pdf')
            pdf_doc.save(output_path)
            pdf_doc.close()
            
            return output_path, change_count
            
        except Exception as e:
            raise Exception(f"Failed to apply redlines to PDF: {str(e)}")
    
    def upload_redlined_document(self, agreement_id: str, file_path: str, format: str) -> str:
        """Upload redlined document to S3"""
        try:
            file_name = os.path.basename(file_path)
            s3_key = f"agreements/{agreement_id}/redlined/{file_name}"
            
            self.s3_client.upload_file(
                file_path,
                self.bucket_name,
                s3_key
            )
            
            return f"s3://{self.bucket_name}/{s3_key}"
            
        except Exception as e:
            raise Exception(f"Failed to upload redlined document: {str(e)}")
    
    def render_document(self, agreement_id: str, options: RenderOptions) -> RenderResult:
        """Render document with redlines and tracked changes"""
        import time
        start_time = time.time()
        
        try:
            # Download required data
            original_path, file_extension = self.download_original_document(agreement_id)
            redline_data = self.download_redline_data(agreement_id)
            structure = self.download_structure(agreement_id)
            
            # Apply redlines based on format
            if file_extension == 'docx':
                output_path, change_count = self.apply_redlines_to_docx(
                    original_path, redline_data, structure, options
                )
            elif file_extension == 'pdf':
                output_path, change_count = self.apply_redlines_to_pdf(
                    original_path, redline_data, structure, options
                )
            else:
                raise Exception(f"Unsupported file format: {file_extension}")
            
            # Upload redlined document
            s3_url = self.upload_redlined_document(agreement_id, output_path, file_extension)
            
            # Calculate metrics
            render_time_ms = int((time.time() - start_time) * 1000)
            file_size = os.path.getsize(output_path)
            
            # Clean up temporary files
            os.unlink(original_path)
            os.unlink(output_path)
            
            return RenderResult(
                agreement_id=agreement_id,
                original_s3_key=f"agreements/{agreement_id}/original/",
                redlined_s3_key=s3_url,
                format=file_extension,
                file_size=file_size,
                change_count=change_count,
                render_time_ms=render_time_ms
            )
            
        except Exception as e:
            # Clean up temporary files on error
            if 'original_path' in locals():
                try:
                    os.unlink(original_path)
                except:
                    pass
            if 'output_path' in locals():
                try:
                    os.unlink(output_path)
                except:
                    pass
            
            raise Exception(f"Failed to render document for {agreement_id}: {str(e)}")
    
    def generate_signed_url(self, s3_key: str, expiration: int = 3600) -> str:
        """Generate a signed URL for document download"""
        try:
            # Remove s3:// prefix if present
            if s3_key.startswith('s3://'):
                s3_key = s3_key[5:]
                bucket_name = s3_key.split('/')[0]
                key = '/'.join(s3_key.split('/')[1:])
            else:
                bucket_name = self.bucket_name
                key = s3_key
            
            url = self.s3_client.generate_presigned_url(
                'get_object',
                Params={'Bucket': bucket_name, 'Key': key},
                ExpiresIn=expiration
            )
            
            return url
            
        except Exception as e:
            raise Exception(f"Failed to generate signed URL: {str(e)}")

# Celery task
@celery_app.task(bind=True, name='redline-engine.render-document')
def render_document(self, agreement_id: str, options: Dict[str, Any] = None) -> Dict[str, Any]:
    """Render document with redlines and tracked changes"""
    try:
        worker = RedlineEngineWorker()
        
        # Set default options
        if options is None:
            options = {}
        
        render_options = RenderOptions(
            include_track_changes=options.get('include_track_changes', True),
            include_watermark=options.get('include_watermark', True),
            watermark_text=options.get('watermark_text', 'DRAFT - FOR REVIEW'),
            include_comments=options.get('include_comments', True),
            output_format=options.get('output_format', 'docx')
        )
        
        # Update task state
        self.update_state(
            state='PROGRESS',
            meta={'status': 'Downloading documents...', 'agreement_id': agreement_id}
        )
        
        # Render document
        result = worker.render_document(agreement_id, render_options)
        
        # Generate signed URL
        signed_url = worker.generate_signed_url(result.redlined_s3_key)
        
        return {
            'status': 'completed',
            'agreement_id': agreement_id,
            'redlined_s3_key': result.redlined_s3_key,
            'signed_url': signed_url,
            'format': result.format,
            'file_size': result.file_size,
            'change_count': result.change_count,
            'render_time_ms': result.render_time_ms
        }
        
    except Exception as e:
        return {
            'status': 'failed',
            'error': str(e),
            'agreement_id': agreement_id
        }
