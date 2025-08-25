# Created automatically by Cursor AI (2024-12-19)
from celery import shared_task
import structlog
import boto3
import hashlib
import os
from typing import Dict, Any, Optional
from pathlib import Path
import tempfile

# Document processing imports
import mammoth
from PyPDF2 import PdfReader
import fitz  # PyMuPDF
import docx
from PIL import Image
import pytesseract

logger = structlog.get_logger()

class DocumentIngestWorker:
    def __init__(self):
        self.s3_client = boto3.client(
            's3',
            endpoint_url=os.getenv('S3_ENDPOINT_URL'),
            aws_access_key_id=os.getenv('S3_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('S3_SECRET_ACCESS_KEY'),
            region_name=os.getenv('S3_REGION', 'us-east-1')
        )
        self.bucket_name = os.getenv('S3_BUCKET_NAME', 'contract-intelligence')

    def download_file(self, s3_key: str) -> str:
        """Download file from S3 to temporary location."""
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=Path(s3_key).suffix)
        try:
            self.s3_client.download_file(self.bucket_name, s3_key, temp_file.name)
            return temp_file.name
        except Exception as e:
            logger.error("Failed to download file from S3", s3_key=s3_key, error=str(e))
            raise

    def upload_file(self, local_path: str, s3_key: str) -> str:
        """Upload file to S3."""
        try:
            self.s3_client.upload_file(local_path, self.bucket_name, s3_key)
            return f"s3://{self.bucket_name}/{s3_key}"
        except Exception as e:
            logger.error("Failed to upload file to S3", s3_key=s3_key, error=str(e))
            raise

    def extract_text_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text and structure from DOCX file."""
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs with their formatting
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name,
                        'runs': [{'text': run.text, 'bold': run.bold, 'italic': run.italic} for run in para.runs]
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Extract headers and footers
            headers = []
            footers = []
            for section in doc.sections:
                if section.header.paragraphs:
                    headers.extend([p.text for p in section.header.paragraphs if p.text.strip()])
                if section.footer.paragraphs:
                    footers.extend([p.text for p in section.footer.paragraphs if p.text.strip()])
            
            return {
                'type': 'docx',
                'paragraphs': paragraphs,
                'tables': tables,
                'headers': headers,
                'footers': footers,
                'success': True
            }
            
        except Exception as e:
            logger.error("Failed to extract text from DOCX", file_path=file_path, error=str(e))
            return {'type': 'docx', 'success': False, 'error': str(e)}

    def extract_text_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF with OCR fallback."""
        try:
            # Try text extraction first
            doc = fitz.open(file_path)
            text_content = []
            images = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                text = page.get_text()
                if text.strip():
                    text_content.append({
                        'page': page_num + 1,
                        'text': text,
                        'blocks': page.get_text("dict")["blocks"]
                    })
                else:
                    # No text found, try OCR
                    pix = page.get_pixmap()
                    img_data = pix.tobytes("png")
                    images.append({
                        'page': page_num + 1,
                        'image_data': img_data
                    })
            
            doc.close()
            
            # Perform OCR on images if needed
            ocr_results = []
            if images:
                logger.info("Performing OCR on pages without text", pages=[img['page'] for img in images])
                for img_data in images:
                    # Save image temporarily for OCR
                    temp_img = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                    temp_img.write(img_data['image_data'])
                    temp_img.close()
                    
                    try:
                        # Perform OCR
                        image = Image.open(temp_img.name)
                        ocr_text = pytesseract.image_to_string(image)
                        ocr_results.append({
                            'page': img_data['page'],
                            'text': ocr_text
                        })
                        os.unlink(temp_img.name)
                    except Exception as e:
                        logger.error("OCR failed for page", page=img_data['page'], error=str(e))
                        os.unlink(temp_img.name)
            
            return {
                'type': 'pdf',
                'text_content': text_content,
                'ocr_results': ocr_results,
                'success': True
            }
            
        except Exception as e:
            logger.error("Failed to extract text from PDF", file_path=file_path, error=str(e))
            return {'type': 'pdf', 'success': False, 'error': str(e)}

    def normalize_document(self, extracted_content: Dict[str, Any]) -> Dict[str, Any]:
        """Normalize extracted content into standard format."""
        normalized = {
            'document_type': extracted_content['type'],
            'sections': [],
            'metadata': {
                'total_pages': 0,
                'has_tables': False,
                'has_images': False,
                'ocr_used': False
            }
        }
        
        if extracted_content['type'] == 'docx':
            # Process DOCX content
            current_section = {'heading': '', 'content': [], 'level': 0}
            
            for para in extracted_content['paragraphs']:
                text = para['text'].strip()
                style = para['style'].lower()
                
                # Detect headings
                if any(keyword in style for keyword in ['heading', 'title', 'header']):
                    if current_section['content']:
                        normalized['sections'].append(current_section)
                    
                    current_section = {
                        'heading': text,
                        'content': [],
                        'level': self._get_heading_level(style)
                    }
                else:
                    current_section['content'].append(text)
            
            # Add last section
            if current_section['content']:
                normalized['sections'].append(current_section)
            
            # Add metadata
            normalized['metadata']['has_tables'] = len(extracted_content.get('tables', [])) > 0
            
        elif extracted_content['type'] == 'pdf':
            # Process PDF content
            for page_content in extracted_content['text_content']:
                page_text = page_content['text']
                # Simple section detection based on font size and formatting
                sections = self._detect_sections_from_pdf(page_text, page_content['blocks'])
                normalized['sections'].extend(sections)
            
            # Add OCR results
            if extracted_content.get('ocr_results'):
                normalized['metadata']['ocr_used'] = True
                for ocr_result in extracted_content['ocr_results']:
                    # Add OCR text as additional sections
                    normalized['sections'].append({
                        'heading': f'OCR Page {ocr_result["page"]}',
                        'content': [ocr_result['text']],
                        'level': 1,
                        'ocr': True
                    })
            
            normalized['metadata']['total_pages'] = len(extracted_content['text_content'])
        
        return normalized

    def _get_heading_level(self, style: str) -> int:
        """Extract heading level from style name."""
        if 'heading 1' in style or 'title' in style:
            return 1
        elif 'heading 2' in style:
            return 2
        elif 'heading 3' in style:
            return 3
        else:
            return 1

    def _detect_sections_from_pdf(self, text: str, blocks: list) -> list:
        """Detect sections from PDF text and formatting blocks."""
        sections = []
        lines = text.split('\n')
        current_section = {'heading': '', 'content': [], 'level': 1}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Simple heuristic: lines with all caps or ending with numbers might be headings
            if (line.isupper() or 
                line.endswith('.') and any(char.isdigit() for char in line) or
                len(line) < 100 and line.endswith(':')):
                
                if current_section['content']:
                    sections.append(current_section)
                
                current_section = {
                    'heading': line,
                    'content': [],
                    'level': 1
                }
            else:
                current_section['content'].append(line)
        
        if current_section['content']:
            sections.append(current_section)
        
        return sections

    def calculate_sha256(self, file_path: str) -> str:
        """Calculate SHA256 hash of file."""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                sha256_hash.update(chunk)
        return sha256_hash.hexdigest()

@shared_task(bind=True)
def ingest_document(self, file_id: str, agreement_id: str):
    """Ingest and parse document."""
    logger.info("Starting document ingestion", file_id=file_id, agreement_id=agreement_id)
    
    worker = DocumentIngestWorker()
    
    try:
        # TODO: Get file metadata from database
        # For now, assume we have the S3 key
        s3_key = f"uploads/{file_id}"
        
        # Download file from S3
        local_file_path = worker.download_file(s3_key)
        
        # Calculate file hash
        file_hash = worker.calculate_sha256(local_file_path)
        
        # Extract text based on file type
        file_extension = Path(local_file_path).suffix.lower()
        
        if file_extension == '.docx':
            extracted_content = worker.extract_text_from_docx(local_file_path)
        elif file_extension == '.pdf':
            extracted_content = worker.extract_text_from_pdf(local_file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not extracted_content['success']:
            raise Exception(f"Failed to extract content: {extracted_content.get('error')}")
        
        # Normalize content
        normalized_content = worker.normalize_document(extracted_content)
        
        # Upload normalized content to S3
        normalized_key = f"processed/{file_id}/normalized.json"
        import json
        temp_json = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json')
        json.dump(normalized_content, temp_json)
        temp_json.close()
        
        worker.upload_file(temp_json.name, normalized_key)
        os.unlink(temp_json.name)
        
        # Cleanup
        os.unlink(local_file_path)
        
        # TODO: Update database with processing results
        # TODO: Trigger next step in workflow (structure parsing)
        
        logger.info("Document ingestion completed", 
                   file_id=file_id, 
                   agreement_id=agreement_id,
                   sections_count=len(normalized_content['sections']))
        
        return {
            "status": "success", 
            "file_id": file_id,
            "sections_count": len(normalized_content['sections']),
            "metadata": normalized_content['metadata']
        }
        
    except Exception as e:
        logger.error("Document ingestion failed", file_id=file_id, error=str(e))
        raise
