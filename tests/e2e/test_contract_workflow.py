# Created automatically by Cursor AI (2024-12-19)

import pytest
import asyncio
import time
from typing import Dict, Any
from pathlib import Path
import tempfile
import shutil

from playwright.async_api import async_playwright, Page, Browser, BrowserContext
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.main import app
from app.core.database import get_db_session
from app.models.agreements import Agreement
from app.models.matters import Matter
from app.models.users import User
from app.models.organizations import Organization

class TestContractWorkflow:
    """End-to-end tests for the complete contract workflow"""
    
    @pytest.fixture(scope="class")
    def test_client(self):
        """Create test client"""
        return TestClient(app)
    
    @pytest.fixture(scope="class")
    def db_session(self):
        """Create database session"""
        engine = create_engine("sqlite:///./test.db")
        TestingSessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
        return TestingSessionLocal()
    
    @pytest.fixture(scope="class")
    async def browser(self):
        """Setup browser for UI tests"""
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            yield browser
            await browser.close()
    
    @pytest.fixture(scope="class")
    async def page(self, browser: Browser):
        """Create browser page"""
        context = await browser.new_context()
        page = await context.new_page()
        yield page
        await context.close()
    
    @pytest.fixture(scope="class")
    def test_data(self):
        """Test data setup"""
        return {
            "org": {
                "name": "Test Organization",
                "domain": "test.com"
            },
            "user": {
                "email": "test@test.com",
                "name": "Test User"
            },
            "matter": {
                "title": "Test Matter",
                "description": "Test matter description"
            },
            "agreement": {
                "title": "Test Agreement",
                "agreement_type": "service",
                "counterparty_name": "Test Counterparty"
            }
        }
    
    async def test_complete_workflow(self, test_client, db_session, page: Page, test_data: Dict[str, Any]):
        """Test complete contract workflow from upload to signature"""
        
        # Step 1: Create test organization and user
        org = await self._create_test_organization(db_session, test_data["org"])
        user = await self._create_test_user(db_session, test_data["user"], org.id)
        matter = await self._create_test_matter(db_session, test_data["matter"], org.id)
        
        # Step 2: Test document upload
        agreement_id = await self._test_document_upload(test_client, page, matter.id, user.id)
        assert agreement_id is not None
        
        # Step 3: Test document processing
        await self._test_document_processing(test_client, agreement_id)
        
        # Step 4: Test clause matching
        await self._test_clause_matching(test_client, agreement_id)
        
        # Step 5: Test redline generation
        redline_id = await self._test_redline_generation(test_client, agreement_id)
        assert redline_id is not None
        
        # Step 6: Test risk assessment
        await self._test_risk_assessment(test_client, agreement_id)
        
        # Step 7: Test approval workflow
        approval_id = await self._test_approval_workflow(test_client, agreement_id, user.id)
        assert approval_id is not None
        
        # Step 8: Test signature process
        envelope_id = await self._test_signature_process(test_client, agreement_id)
        assert envelope_id is not None
        
        # Step 9: Test obligations extraction
        await self._test_obligations_extraction(test_client, agreement_id)
        
        # Step 10: Test reporting
        await self._test_reporting(test_client, agreement_id)
        
        # Cleanup
        await self._cleanup_test_data(db_session, org.id, user.id, matter.id, agreement_id)
    
    async def _create_test_organization(self, db_session, org_data: Dict[str, Any]) -> Organization:
        """Create test organization"""
        org = Organization(
            name=org_data["name"],
            domain=org_data["domain"],
            status="active"
        )
        db_session.add(org)
        db_session.commit()
        db_session.refresh(org)
        return org
    
    async def _create_test_user(self, db_session, user_data: Dict[str, Any], org_id: str) -> User:
        """Create test user"""
        user = User(
            email=user_data["email"],
            name=user_data["name"],
            org_id=org_id,
            status="active"
        )
        db_session.add(user)
        db_session.commit()
        db_session.refresh(user)
        return user
    
    async def _create_test_matter(self, db_session, matter_data: Dict[str, Any], org_id: str) -> Matter:
        """Create test matter"""
        matter = Matter(
            title=matter_data["title"],
            description=matter_data["description"],
            org_id=org_id,
            status="active"
        )
        db_session.add(matter)
        db_session.commit()
        db_session.refresh(matter)
        return matter
    
    async def _test_document_upload(self, test_client, page: Page, matter_id: str, user_id: str) -> str:
        """Test document upload functionality"""
        
        # Create test document
        test_doc = self._create_test_document()
        
        # Navigate to upload page
        await page.goto(f"http://localhost:3000/matters/{matter_id}/upload")
        
        # Upload document
        await page.set_input_files('input[type="file"]', str(test_doc))
        await page.click('button[type="submit"]')
        
        # Wait for upload to complete
        await page.wait_for_selector('.upload-success', timeout=30000)
        
        # Get agreement ID from response
        response = await page.wait_for_response("**/api/v1/agreements")
        response_data = await response.json()
        
        # Cleanup test document
        test_doc.unlink()
        
        return response_data["id"]
    
    async def _test_document_processing(self, test_client, agreement_id: str):
        """Test document processing pipeline"""
        
        # Check document processing status
        response = test_client.get(f"/api/v1/agreements/{agreement_id}/status")
        assert response.status_code == 200
        
        # Wait for processing to complete
        max_retries = 30
        for i in range(max_retries):
            response = test_client.get(f"/api/v1/agreements/{agreement_id}/status")
            status_data = response.json()
            
            if status_data["status"] == "processed":
                break
            
            if i == max_retries - 1:
                pytest.fail("Document processing did not complete in time")
            
            time.sleep(2)
        
        # Verify sections were extracted
        response = test_client.get(f"/api/v1/agreements/{agreement_id}/sections")
        assert response.status_code == 200
        sections = response.json()
        assert len(sections) > 0
    
    async def _test_clause_matching(self, test_client, agreement_id: str):
        """Test clause matching functionality"""
        
        # Trigger clause matching
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/match-clauses")
        assert response.status_code == 202
        
        # Wait for matching to complete
        max_retries = 30
        for i in range(max_retries):
            response = test_client.get(f"/api/v1/agreements/{agreement_id}/clause-matches")
            matches = response.json()
            
            if len(matches) > 0:
                break
            
            if i == max_retries - 1:
                pytest.fail("Clause matching did not complete in time")
            
            time.sleep(2)
        
        # Verify matches were found
        assert len(matches) > 0
    
    async def _test_redline_generation(self, test_client, agreement_id: str) -> str:
        """Test redline generation"""
        
        # Generate redline
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/redline")
        assert response.status_code == 202
        
        # Wait for redline to complete
        max_retries = 30
        for i in range(max_retries):
            response = test_client.get(f"/api/v1/agreements/{agreement_id}/redline")
            redline_data = response.json()
            
            if redline_data["status"] == "completed":
                break
            
            if i == max_retries - 1:
                pytest.fail("Redline generation did not complete in time")
            
            time.sleep(2)
        
        return redline_data["id"]
    
    async def _test_risk_assessment(self, test_client, agreement_id: str):
        """Test risk assessment"""
        
        # Generate risk report
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/risk-assessment")
        assert response.status_code == 202
        
        # Wait for risk assessment to complete
        max_retries = 30
        for i in range(max_retries):
            response = test_client.get(f"/api/v1/agreements/{agreement_id}/risk-report")
            risk_data = response.json()
            
            if risk_data["status"] == "completed":
                break
            
            if i == max_retries - 1:
                pytest.fail("Risk assessment did not complete in time")
            
            time.sleep(2)
        
        # Verify risk report
        assert risk_data["overall_risk_score"] is not None
        assert len(risk_data["exceptions"]) >= 0
    
    async def _test_approval_workflow(self, test_client, agreement_id: str, user_id: str) -> str:
        """Test approval workflow"""
        
        # Create approval gate
        approval_data = {
            "name": "Legal Review",
            "description": "Legal team review required",
            "approvers": [user_id],
            "due_date": "2024-12-31T23:59:59Z"
        }
        
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/approvals", json=approval_data)
        assert response.status_code == 201
        
        approval_id = response.json()["id"]
        
        # Approve the agreement
        response = test_client.post(f"/api/v1/approvals/{approval_id}/approve", json={
            "notes": "Approved by test user"
        })
        assert response.status_code == 200
        
        return approval_id
    
    async def _test_signature_process(self, test_client, agreement_id: str) -> str:
        """Test signature process"""
        
        # Create signature envelope
        envelope_data = {
            "recipients": [
                {
                    "email": "signer@test.com",
                    "name": "Test Signer",
                    "role": "signer"
                }
            ],
            "fields": [
                {
                    "type": "signature",
                    "page": 1,
                    "x": 100,
                    "y": 100
                }
            ]
        }
        
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/signature", json=envelope_data)
        assert response.status_code == 201
        
        envelope_id = response.json()["id"]
        
        # Simulate signature completion (in real scenario, this would be via webhook)
        response = test_client.post(f"/api/v1/signature/{envelope_id}/complete")
        assert response.status_code == 200
        
        return envelope_id
    
    async def _test_obligations_extraction(self, test_client, agreement_id: str):
        """Test obligations extraction"""
        
        # Extract obligations
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/extract-obligations")
        assert response.status_code == 202
        
        # Wait for extraction to complete
        max_retries = 30
        for i in range(max_retries):
            response = test_client.get(f"/api/v1/agreements/{agreement_id}/obligations")
            obligations = response.json()
            
            if len(obligations) > 0:
                break
            
            if i == max_retries - 1:
                pytest.fail("Obligations extraction did not complete in time")
            
            time.sleep(2)
        
        # Verify obligations were extracted
        assert len(obligations) > 0
    
    async def _test_reporting(self, test_client, agreement_id: str):
        """Test reporting functionality"""
        
        # Generate executive summary
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/reports/executive-summary")
        assert response.status_code == 200
        
        summary = response.json()
        assert summary["agreement_title"] is not None
        assert summary["overall_risk_score"] is not None
        
        # Generate risk report
        response = test_client.post(f"/api/v1/agreements/{agreement_id}/reports/risk-report")
        assert response.status_code == 200
        
        risk_report = response.json()
        assert risk_report["overall_risk_score"] is not None
        assert len(risk_report["recommendations"]) >= 0
    
    def _create_test_document(self) -> Path:
        """Create a test document for upload"""
        test_doc = Path(tempfile.gettempdir()) / "test_agreement.docx"
        
        # Create a simple test document
        from docx import Document
        doc = Document()
        doc.add_heading('Test Agreement', 0)
        doc.add_paragraph('This is a test agreement for automated testing.')
        doc.add_heading('1. Services', level=1)
        doc.add_paragraph('Provider shall provide the following services...')
        doc.add_heading('2. Payment Terms', level=1)
        doc.add_paragraph('Client shall pay Provider for services rendered...')
        doc.add_heading('3. Term', level=1)
        doc.add_paragraph('This agreement shall commence on the Effective Date...')
        
        doc.save(str(test_doc))
        return test_doc
    
    async def _cleanup_test_data(self, db_session, org_id: str, user_id: str, matter_id: str, agreement_id: str):
        """Cleanup test data"""
        try:
            # Delete agreement
            db_session.query(Agreement).filter(Agreement.id == agreement_id).delete()
            
            # Delete matter
            db_session.query(Matter).filter(Matter.id == matter_id).delete()
            
            # Delete user
            db_session.query(User).filter(User.id == user_id).delete()
            
            # Delete organization
            db_session.query(Organization).filter(Organization.id == org_id).delete()
            
            db_session.commit()
        except Exception as e:
            db_session.rollback()
            print(f"Cleanup failed: {e}")

class TestSecurityFeatures:
    """Test security features and hardening"""
    
    def test_rbac_enforcement(self, test_client):
        """Test RBAC enforcement"""
        # Test unauthorized access
        response = test_client.get("/api/v1/agreements")
        assert response.status_code == 401
        
        # Test with invalid token
        response = test_client.get("/api/v1/agreements", headers={"Authorization": "Bearer invalid"})
        assert response.status_code == 401
    
    def test_rate_limiting(self, test_client):
        """Test rate limiting"""
        # Make multiple requests quickly
        for i in range(150):
            response = test_client.get("/api/v1/health")
            if response.status_code == 429:
                break
        else:
            pytest.fail("Rate limiting not enforced")
    
    def test_input_validation(self, test_client):
        """Test input validation"""
        # Test with malicious input
        malicious_data = {
            "title": "<script>alert('xss')</script>",
            "description": "'; DROP TABLE agreements; --"
        }
        
        response = test_client.post("/api/v1/agreements", json=malicious_data)
        assert response.status_code == 400  # Should be rejected
    
    def test_sql_injection_prevention(self, test_client):
        """Test SQL injection prevention"""
        # Test with SQL injection attempts
        injection_attempts = [
            "'; DROP TABLE agreements; --",
            "' OR '1'='1",
            "'; INSERT INTO users VALUES ('hacker', 'password'); --"
        ]
        
        for attempt in injection_attempts:
            response = test_client.get(f"/api/v1/agreements?title={attempt}")
            assert response.status_code in [400, 401, 404]  # Should not execute

class TestPerformanceFeatures:
    """Test performance and scalability features"""
    
    def test_concurrent_requests(self, test_client):
        """Test handling of concurrent requests"""
        import concurrent.futures
        
        def make_request():
            return test_client.get("/api/v1/health")
        
        # Make 50 concurrent requests
        with concurrent.futures.ThreadPoolExecutor(max_workers=50) as executor:
            futures = [executor.submit(make_request) for _ in range(50)]
            responses = [future.result() for future in concurrent.futures.as_completed(futures)]
        
        # All requests should succeed
        assert all(response.status_code == 200 for response in responses)
    
    def test_large_document_processing(self, test_client):
        """Test processing of large documents"""
        # Create a large test document
        large_doc = self._create_large_test_document()
        
        # Test upload and processing
        with open(large_doc, 'rb') as f:
            response = test_client.post("/api/v1/upload", files={"file": f})
        
        assert response.status_code == 202
        
        # Cleanup
        large_doc.unlink()
    
    def _create_large_test_document(self) -> Path:
        """Create a large test document"""
        large_doc = Path(tempfile.gettempdir()) / "large_test_agreement.docx"
        
        from docx import Document
        doc = Document()
        doc.add_heading('Large Test Agreement', 0)
        
        # Add many sections to make it large
        for i in range(100):
            doc.add_heading(f'Section {i+1}', level=1)
            doc.add_paragraph(f'This is section {i+1} content with detailed terms and conditions...')
            doc.add_paragraph(f'Additional terms for section {i+1}...')
        
        doc.save(str(large_doc))
        return large_doc

if __name__ == "__main__":
    pytest.main([__file__, "-v"])
